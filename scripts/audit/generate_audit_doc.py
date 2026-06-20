#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génère le document Word d'audit pipeline de la donnée
LyonFlowFull — 2026-06-20

Style: ModernCorporate — français
Source: 4 audits détaillés (Bronze, Silver, Gold, ML+Routing)
Format: docx via python-docx 1.2.0
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# Couleurs
C_TITLE = RGBColor(0x1F, 0x36, 0x4D)         # bleu nuit corporate
C_H1 = RGBColor(0x1F, 0x36, 0x4D)
C_H2 = RGBColor(0x2E, 0x5A, 0x88)
C_H3 = RGBColor(0x3F, 0x6E, 0xA6)
C_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
C_MUTED = RGBColor(0x6A, 0x73, 0x7B)
C_CRITIQUE = RGBColor(0xC0, 0x39, 0x2B)      # rouge
C_MAJEURE = RGBColor(0xE6, 0x7E, 0x22)       # orange
C_MINEURE = RGBColor(0xF3, 0x9C, 0x12)       # ambre
C_COSMETIQUE = RGBColor(0x27, 0xAE, 0x60)    # vert
C_INFO = RGBColor(0x2E, 0x5A, 0x88)          # bleu


def add_shading(cell, color_hex):
    """Ajoute un fond coloré à une cellule."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_borders(cell):
    """Bordures fines grises sur une cellule."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'B0BEC5')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = C_H1
    r.font.name = 'Calibri'
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = C_H2
    r.font.name = 'Calibri'
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = C_H3
    r.font.name = 'Calibri'
    return p


def add_para(doc, text, italic=False, color=None, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color or C_TEXT
    r.italic = italic
    r.font.name = 'Calibri'
    return p


def add_mono(doc, text):
    """Code/file paths en monospace."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.name = 'Consolas'
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p


def add_badge(doc, level, text):
    """Badge coloré en début de paragraphe."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    colors = {
        'CRITIQUE': (C_CRITIQUE, 'FFE5E0'),
        'MAJEURE': (C_MAJEURE, 'FDEBD0'),
        'MINEURE': (C_MINEURE, 'FEF5E7'),
        'COSMETIQUE': (C_COSMETIQUE, 'D5F5E3'),
        'INFO': (C_INFO, 'D6EAF8'),
    }
    fg, bg = colors[level]
    r = p.add_run(f'  {level}  ')
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = fg
    rpr = r._element
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rpr.insert(0, rFonts)

    r2 = p.add_run(f'  {text}')
    r2.font.size = Pt(10.5)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = C_TEXT
    return p


def add_table(doc, headers, rows, col_widths=None, header_bg='1F364D'):
    """Tableau avec en-tête coloré et cellules alignées."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = 'Calibri'
        add_shading(cell, header_bg)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = 'Calibri'
            r.font.color.rgb = C_TEXT
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if ri % 2 == 1:
                add_shading(cell, 'F4F6F8')
    return table


# =========================================================
# Construction du document
# =========================================================

doc = Document()

# Page setup
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = C_TEXT


# =========================================================
# PAGE DE GARDE
# =========================================================

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('LYONFLOWFULL')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = C_MUTED
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run('Plateforme MLOps — Trafic multimodal Lyon')
r.font.size = Pt(10)
r.font.color.rgb = C_MUTED
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
r = p.add_run('Audit du pipeline de la donnée')
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = C_TITLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
r = p.add_run('Couches Bronze · Silver · Gold · Modèles ML · Routing')
r.font.size = Pt(14)
r.font.color.rgb = C_H2
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(180)
r = p.add_run('Version du document')
r.font.size = Pt(9)
r.font.color.rgb = C_MUTED
r.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('1.0 — 2026-06-20')
r.font.size = Pt(11)
r.font.color.rgb = C_TEXT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
r = p.add_run('Préparé pour')
r.font.size = Pt(9)
r.font.color.rgb = C_MUTED
r.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Patrice DUCLOS — Senior Data Analyst / Architecte en IA')
r.font.size = Pt(11)
r.font.color.rgb = C_TEXT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
r = p.add_run('Méthodologie')
r.font.size = Pt(9)
r.font.color.rgb = C_MUTED
r.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Lecture exhaustive du code source · citation fichier:ligne · aucun mock')
r.font.size = Pt(10.5)
r.font.color.rgb = C_TEXT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
r = p.add_run('Périmètre analysé')
r.font.size = Pt(9)
r.font.color.rgb = C_MUTED
r.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('~170 fichiers Python · 13 DAGs Airflow · 22 migrations SQL · 3 modèles ML')
r.font.size = Pt(10.5)
r.font.color.rgb = C_TEXT

page_break(doc)

# =========================================================
# RÉSUMÉ EXÉCUTIF
# =========================================================

add_h1(doc, 'Résumé exécutif')

add_para(doc,
    "Cet audit couvre l'ensemble du pipeline de la donnée du projet LyonFlowFull à la date du 2026-06-20. "
    "Il s'appuie exclusivement sur la lecture du code source (Python, SQL, YAML, configurations Airflow) et "
    "des migrations SQL officielles. Aucune donnée fictive, aucun mock, aucune extrapolation."
)

add_para(doc,
    "Le projet suit une architecture Medallion (Bronze → Silver → Gold) avec 8 collecteurs Bronze, 5 tables Silver "
    "de nettoyage, ~15 tables Gold actives (dont 4 vues matérialisées stratégiques) et 3 modèles ML (XGBoost Speed, "
    "XGBoost Vélov, ST-GRU-GNN spatio-temporel). Le dashboard Streamlit sert 3 personas (Usager, Pro TCL, Élu) "
    "via 18 pages et 51 widgets."
)

add_h2(doc, 'Verdict global')

add_table(doc,
    headers=['Couche', 'État fonctionnel', 'Risque principal'],
    rows=[
        ['Bronze — collecte', '⚠ Partiellement opérationnel',
         'INSERT silencieusement cassé sur calendrier_scolaire + jours_feries (run() swallow les exceptions)'],
        ['Silver — nettoyage', '❌ Cassé en local sur plusieurs tables',
         'Schéma DDL ≠ code INSERT pour 3 tables (trafic_boucles_clean, meteo_hourly, chantiers_actifs)'],
        ['Gold — analytique', '⚠ Partiellement opérationnel',
         '6 tables lues par le code mais jamais peuplées (velov_predictions, channels_ref, etc.)'],
        ['Modèles ML', 'XGBoost prod ✅ · Vélov dormant · GNN inactif',
         'XGBoost Vélov : features ne matchent pas la table → KeyError si activé'],
        ['Routing multimodal', 'Fonctionnel en local',
         'Horizon H+1h non utilisé dans le pathfinder (vitesses H+0 toujours)'],
    ],
    col_widths=[Cm(4.5), Cm(5.5), Cm(7)],
)

add_h2(doc, 'Chiffres-clés de l\'audit')

add_table(doc,
    headers=['Indicateur', 'Valeur'],
    rows=[
        ['Fichiers Python audités', '~170 (~22 000 lignes)'],
        ['DAGs Airflow', '13 (10 actifs + 1 cron + 1 archive + 1 TomTom)'],
        ['Migrations SQL', '22 (init-db + 21 incréments)'],
        ['Collecteurs Bronze actifs', '9 (dont 1 ajouté Sprint 13+ : TomTom)'],
        ['Tables Bronze', '17 (dont 6 orphelines)'],
        ['Tables Silver', '5'],
        ['Tables/Vues Gold actives', '~25 (dont 4 vues matérialisées)'],
        ['Tables Gold fantômes lues par le code', '6'],
        ['Doutes identifiés', '~80 (12 critiques, 30 majeures, 38 mineures/cosmétiques)'],
        ['Tests verts (au 2026-06-19)', '301 / 4 SKIP / 14 deselected'],
    ],
    col_widths=[Cm(7), Cm(10)],
)

add_h2(doc, 'Top 5 — ce qui doit être corrigé en priorité')

p = doc.add_paragraph()
items = [
    ('1.', 'CRITIQUE', 'CalendrierScolaire + JoursFeries — INSERT cassés (colonnes fetched_at/raw_data absentes du DDL) + run() swallow les exceptions. Conséquence : calendriers vides, _is_vacances/_is_ferie crashent à l\'usage.'),
    ('2.', 'CRITIQUE', 'silver.trafic_boucles_clean — schéma DDL (speed_kmh) ≠ code INSERT (vitesse_kmh), contrainte UNIQUE absente du DDL. Si DDL non altéré en prod, ON CONFLICT lève NoUniqueOrPrimaryKeyError à chaque exécution.'),
    ('3.', 'CRITIQUE', 'silver.meteo_hourly — schema DDL (temperature_2m, precipitation) ≠ code INSERT (temperature_c, rain_mm). Dette documentée 2 fois (CHANGELOG, AUDIT_PIPELINE_2026-06-12), patch in-place sur VPS non commité.'),
    ('4.', 'CRITIQUE', 'gold.velov_predictions, gold.amenagements_history, gold.fact_correlation_matrix, gold.mv_kpis_12_months — lues par le code applicatif mais jamais peuplées. Widgets crashent silencieusement via _df_from_query:130 qui retourne un DataFrame vide.'),
    ('5.', 'CRITIQUE', 'XGBoost Vélov — features attendues (bikes_lag_1/2/3, rolling_mean_3h) ≠ colonnes de la table (lag_30min, lag_60min, rolling_mean_1h). Si le DAG hourly :50 est réactivé, KeyError garanti. AGENTS.md le mentionne comme dette technique connue non résolue.'),
]
for num, level, text in items:
    add_badge(doc, level, f'P{num} {text}')

page_break(doc)

# =========================================================
# 1. VUE D'ENSEMBLE
# =========================================================

add_h1(doc, '1. Vue d\'ensemble du pipeline')

add_para(doc,
    "LyonFlowFull implémente une architecture Medallion à trois couches (Bronze, Silver, Gold) "
    "avec une couche de référence (referentiel) et une couche de monitoring (drift, santé, RGPD). "
    "L'orchestration repose sur Apache Airflow 2.9 avec 13 DAGs, et la persistance sur PostgreSQL 16 + PostGIS."
)

add_h2(doc, '1.1 Schéma de flux')

add_para(doc, "Flux nominal d'une donnée collectée jusqu'à sa consommation par un widget :")
add_mono(doc, """
  [Source externe]                                                        [Widget Streamlit]
       │                                                                          ▲
       ▼                                                                          │
  ┌─────────────────┐                                                            │
  │ Bronze — 8 coll.│  9 tables alimentées en continu                          │
  │  (DAG */5min)   │  • trafic_boucles, tcl_vehicles, velov, meteo, etc.      │
  └────────┬────────┘                                                            │
           │                                                                     │
           ▼  (DAG */5min)                                                       │
  ┌─────────────────┐                                                            │
  │ Silver — 5 tabs │  • trafic_boucles_clean                                   │
  │  (parse + dedup)│  • tcl_vehicles_clean, velov_clean, meteo_hourly          │
  │                 │  • chantiers_actifs                                        │
  └────────┬────────┘                                                            │
           │                                                                     │
           ▼  (DAG */10min + quotidien 03h)                                     │
  ┌──────────────────────────────────────────────────────────────────────────┐  │
  │ Gold — features + analytique + vues matérialisées                        │  │
  │  • traffic_features_live (XGBoost features)                              │──┘
  │  • bus_delay_segments, infrastructure_bottlenecks                        │
  │  • mv_multimodal_grid (Sprint 15+ Axe 1)                                  │
  │  • mv_bus_traffic_spatial (Sprint 15+ Axe 3)                              │
  │  • fn_network_health_score (Sprint 15+ Axe 5)                             │
  │  • trafic_predictions (toutes les 15 min)                                 │
  │  • v_coherence_tomtom_vs_grandlyon, v_tomtom_gl_drift                     │
  └────────┬─────────────────────────────────────────────────────────────────┘
           │
           ├──► [ML : XGBoost H+1h, quotidien 03h00] ──► [gold.trafic_predictions] ──► widgets
           │
           ├──► [ML : XGBoost Vélov — dormant] (features cassées)
           │
           ├──► [ML : ST-GRU-GNN — inactif (DAG paused)]
           │
           └──► [Routing : pathfinder_multimodal + eco_calculator] ──► API FastAPI + widgets
""")

add_h2(doc, '1.2 Couverture par couche')

add_table(doc,
    headers=['Couche', 'Tables actives', 'Vues / MVs', 'Producteur', 'Consommateur principal'],
    rows=[
        ['Bronze', '9 (+ 6 orphelines)', '—', '8 collecteurs + DAG collect_bronze */5min + DAG calendriers @monthly + DAG TomTom */15', 'DAG transform_bronze_to_silver'],
        ['Silver', '5', '—', 'DAG transform_bronze_to_silver */5min', 'DAG transform_silver_to_gold + fonctions _is_ferie/_is_vacances'],
        ['Gold', '~15', '4 vues matérialisées + 7 vues + 1 fonction', 'DAG transform_silver_to_gold */10min + DAG ML quotidien', 'Widgets Streamlit + API FastAPI + routing'],
        ['referentiel', '6 tables + 7 vues + 5 fonctions', '—', 'Migrations SQL (seed INSERT ON CONFLICT)', 'Routing + widgets + pathfinder'],
    ],
    col_widths=[Cm(2.5), Cm(3.5), Cm(3), Cm(4), Cm(4)],
)

page_break(doc)

# =========================================================
# 2. COUCHE BRONZE
# =========================================================

add_h1(doc, '2. Couche Bronze — Collecte')

add_para(doc,
    "La couche Bronze contient 9 tables effectivement alimentées par 8 collecteurs (1 collecteur alimente 2 tables "
    "via override) et 6 tables orphelines (créées par init-db.sql mais jamais peuplées). Tous les collecteurs "
    "héritent d'une classe abstraite DataCollector (src/ingestion/base.py) qui définit le cycle fetch_raw() → "
    "validate() → _save_raw()."
)

add_h2(doc, '2.1 Inventaire des 9 collecteurs actifs')

add_table(doc,
    headers=['#', 'Collecteur', 'Source', 'Schedule', 'Table Bronze', 'Volume/jour'],
    rows=[
        ['1', 'TraficGrandLyon', 'WFS Grand Lyon pvotrafic', '*/5 min', 'trafic_boucles', '~316 800 rows'],
        ['2', 'VelovCollector', 'GBFS station_status + information', '*/5 min', 'velov', '~131 904 rows'],
        ['3', 'MeteoOpenMeteo', 'api.open-meteo.com/v1/forecast', '*/5 min (⚠)', 'meteo', '~2 073 rows'],
        ['4', 'AirQualityOpenMeteo', 'air-quality-api.open-meteo.com', '*/5 min (⚠)', 'air_quality', '~2 073 rows'],
        ['5', 'ChantiersGrandLyon', 'WFS pvochantierperturbant', '*/5 min (⚠)', 'chantiers', '~99 360 rows'],
        ['6', 'TclSiriLite', 'data.grandlyon.com/siri-lite/2.0/vehicle-monitoring.json', '*/5 min', 'tcl_vehicles', '~86 400 rows'],
        ['7', 'TomTomTrafficFlow', 'api.tomtom.com flowSegmentData × 12 tuiles', '*/15 min', 'tomtom_traffic', '0–1 152 rows'],
        ['8', 'CalendrierScolaire', 'data.education.gouv.fr Zone A', '@monthly', 'calendrier_scolaire', '~50 rows/mois'],
        ['9', 'JoursFeries', 'calendrier.api.gouv.fr × 2 ans', '@monthly', 'jours_feries', '~22 rows/mois'],
    ],
    col_widths=[Cm(0.8), Cm(2.8), Cm(4.5), Cm(2.2), Cm(2.5), Cm(2.5)],
)

add_para(doc,
    "⚠ Les fréquences indiquées dans les docstrings des collecteurs (meteo.py:5, air_quality.py:5, chantiers.py:5) "
    "annoncent « 1h » ou « 1x/jour » mais le DAG collect_bronze les exécute toutes les 5 min. Cette divergence "
    "entraîne une sur-collecte massive (chantiers : ~100k rows/jour au lieu de ~350 attendus)."
)

add_h2(doc, '2.2 Schéma typique d\'une table Bronze')

add_para(doc, "Pris sur bronze.trafic_boucles (init-db.sql:339-352) :")
add_mono(doc, """
  CREATE TABLE bronze.trafic_boucles (
      id BIGSERIAL PRIMARY KEY,
      fetched_at TIMESTAMPTZ NOT NULL,
      troncon_id TEXT, troncon_name TEXT,
      debit INTEGER, taux_occupation REAL, vitesse REAL,
      raw_data JSONB,
      geom GEOMETRY(Point, 2154),
      geom_4326 GEOMETRY(Point, 4326),
      CONSTRAINT chk_dual_geom CHECK (
          (geom IS NULL AND geom_4326 IS NULL)
          OR (geom IS NOT NULL AND geom_4326 IS NOT NULL)
      )
  );
  -- Index: GIST geom_2154 + GIST geom_4326 + btree (troncon_id, fetched_at)
  -- UNIQUE (troncon_id, fetched_at) WHERE troncon_id IS NOT NULL  -- ligne 2085
""")

add_para(doc,
    "Constat majeur : sur 7 collecteurs (tous sauf TomTom), les colonnes typées (troncon_id, debit, etc.) sont "
    "TOUJOURS NULL car DataCollector._save_raw() insère uniquement (fetched_at, raw_data) (base.py:183-186). "
    "Le parsing est délégué à la transformation Bronze → Silver."
)

add_h2(doc, '2.3 Cycle de collecte — base.py')

add_para(doc, "Pattern template method appliqué par les 9 collecteurs :")
add_mono(doc, """
  def run(self) -> FetchResult:
      start = time.time()
      result = self.fetch_raw()               # peut lever
      result.duration_ms = int((time.time() - start) * 1000)
      if not self.validate(result):           # default: n_records >= 0
          raise CollectorError(...)
      self._save_raw(result)                  # INSERT (fetched_at, raw_data) VALUES (...)
      self.n_requests += 1
      self.last_success_at = result.fetched_at
      return result
      except Exception as e:                  # ⚠ base.py:145-155
          self.n_failures += 1
          self.last_error = str(e)
          return FetchResult(error=str(e))    # ne lève pas
""")

add_h2(doc, '2.4 Doutes et incohérences — couche Bronze')

add_badge(doc, 'CRITIQUE', 'D1 — CalendrierScolaire + JoursFeries : INSERT cassé en silence')
add_para(doc,
    "Les tables bronze.calendrier_scolaire (init-db.sql:108-115) et bronze.jours_feries (init-db.sql:382-385) "
    "n'ont PAS de colonnes fetched_at ni raw_data. Or _save_raw (base.py:183-186) fait INSERT INTO bronze.{table} "
    "(fetched_at, raw_data). Le crash psycopg2 est capturé par le except de run() (base.py:145-155) qui retourne "
    "FetchResult(error=...) sans raise. La tâche Airflow est marquée SUCCEEDED alors qu\'elle a échoué. "
    "Conséquence : calendriers vides, et _is_vacances() / _is_ferie() (silver_to_gold.py:65-99) crashent à l\'usage."
)

add_badge(doc, 'CRITIQUE', 'D7 — run() swallow les exceptions → tâches Airflow "vertes" en erreur')
add_para(doc,
    "base.py:145-155 capture toute exception et retourne FetchResult(error=...) sans raise. Conséquence : la "
    "tâche apparaît verte dans l\'UI mais l\'erreur est silencieuse. Le callback on_failure_callback ne se "
    "déclenche JAMAIS. Seul le DAG collect_tomtom_traffic (collect_tomtom_traffic.py:48-53) contourne ce problème "
    "en checkant manuellement result.error et en levant RuntimeError."
)

add_badge(doc, 'CRITIQUE', 'D15 — gold.v_source_health référence fetched_at sur tables inexistantes')
add_para(doc,
    "La vue gold.v_source_health (migration_021_source_health.sql:27-112) appelle MAX(fetched_at) sur "
    "bronze.calendrier_scolaire et bronze.jours_feries. Ces tables n\'ont pas de fetched_at (cf. D1). "
    "La vue va lever UndefinedColumn à chaque appel. Consommateurs : widget data_quality_badge (Elu_1) et "
    "source_health_monitor (Pro_6)."
)

add_badge(doc, 'MAJEURE', 'D2 — bronze.trafic_vitesse_brute DROP mentionnée mais pas exécutée')
add_para(doc,
    "AGENTS.md indique qu\'un commit ec37d76 a droppé bronze.trafic_vitesse_brute, mais init-db.sql:621-625 la "
    "crée encore. Sur un VPS vierge, rejouer init-db.sql plantera sur CREATE INDEX (lignes 1820, 1827) pointant "
    "vers une table inexistante après le drop. Aucun fichier drop_trafic_vitesse_brute.sql dans scripts/sql/."
)

add_badge(doc, 'MAJEURE', 'D8 — Docstrings mensongères sur les fréquences')
add_para(doc,
    "meteo.py:5 dit « 1h », air_quality.py:5 dit « 1h », chantiers.py:5 dit « 1x/jour ». "
    "Le DAG collect_bronze schedule */5 * * * * (collect_bronze.py:48). "
    "Impact : chantiers pollue Bronze à ~100k rows/jour au lieu de ~350."
)

add_badge(doc, 'MAJEURE', 'D9 — chantiers pollue Bronze à 100k rows/jour')
add_para(doc,
    "chantiers.py WFS maxFeatures=2000 (chantiers.py:49) × 288 cycles/jour = ~99 360 inserts/jour. "
    "AGENTS.md annonce ~350KB GeoJSON/jour (350 chantiers). "
    "Aucun dédup côté Bronze. Purge Bronze 7j (maintenance.py:160) → ~700k rows en base permanente pour des "
    "chantiers qui ne changent presque jamais."
)

add_badge(doc, 'MAJEURE', 'D11 — _validate_table whitelist inclut calendriers_scolaire + jours_feries')
add_para(doc,
    "maintenance.py:19-30 inclut bronze.calendrier_scolaire et bronze.jours_feries dans la purge. "
    "Le DELETE WHERE fetched_at < NOW() - ... explosera aussi (fetched_at inexistant, cf. D1)."
)

add_badge(doc, 'MINEURE', 'D3 — Colonnes extracted non peuplées (sauf TomTom)')
add_para(doc,
    "Les colonnes typées de bronze.air_quality, bronze.chantiers, bronze.tcl_vehicles, bronze.trafic_boucles, "
    "bronze.velov, bronze.meteo sont TOUTES NULL car _save_raw met tout dans raw_data (base.py:183-186). "
    "Seul TomTomTrafficFlow._save_raw (override tomtom_traffic.py:478) fait du columnar. "
    "Conséquence : double parsing JSONB inutile à la transformation Silver."
)

add_badge(doc, 'MINEURE', 'D4 — bronze.jours_feries sans index ni clé primaire')
add_para(doc,
    "init-db.sql:382-385 : juste (date_ferie, nom), aucun index. "
    "_is_ferie() (silver_to_gold.py:65-71) fait SELECT EXISTS → seq-scan complet à chaque appel."
)

add_badge(doc, 'MINEURE', 'D5 — Doublon bronze.tomtom_flow vs bronze.tomtom_traffic')
add_para(doc,
    "init-db.sql:552-564 crée bronze.tomtom_flow (nouveau schéma). "
    "create_tomtom_traffic.sql:14 crée bronze.tomtom_traffic (schéma columnar Sprint 13+). "
    "Aucun code n\'écrit dans tomtom_flow. Risque de confusion dans les futures migrations."
)

add_badge(doc, 'MINEURE', 'D6 — 6 tables Bronze orphelines')
add_para(doc,
    "Créées par init-db.sql mais aucun collecteur n\'y écrit : bronze.chantiers_historique (160), "
    "bronze.chantiers_voirie (218), bronze.comptages (254), bronze.parkings (432), "
    "bronze.prix_carburants (468), bronze.pvotrafic_snapshots (302). Espace disque gaspillé, "
    "confusion pour les futurs devs. À DROP en Sprint 16+."
)

add_badge(doc, 'MINEURE', 'D10 — Pas de métriques Prometheus sur les collecteurs')
add_para(doc,
    "n_requests, n_failures, last_success_at, last_error (base.py:97-101) sont des compteurs d\'instance. "
    "Aucune exposition Prometheus, aucun push XCom systématique. Seul TomTomTrafficFlow push manuellement "
    "XCom (collect_tomtom_traffic.py:63,73)."
)

add_badge(doc, 'MINEURE', 'D14 — TomTomTrafficFlow doublement planifié')
add_para(doc,
    "Présent dans REALTIME_COLLECTORS (__init__.py:30) → collect_bronze toutes les 5 min, ET dans DAG "
    "collect_tomtom_traffic (collect_tomtom_traffic.py:93-101) toutes les 15 min. "
    "Le cache 5min (tomtom_traffic.py:96) atténue mais les 2 tâches écrivent en concurrence dans "
    "bronze.tomtom_traffic (protégé par UNIQUE + ON CONFLICT, mais contre-intuitif)."
)

page_break(doc)

# =========================================================
# 3. COUCHE SILVER
# =========================================================

add_h1(doc, '3. Couche Silver — Nettoyage')

add_para(doc,
    "La couche Silver contient 5 tables alimentées par le DAG transform_bronze_to_silver (DAG id éponyme, "
    "schedule */5min, retries=1). Le code de transformation (src/transformation/bronze_to_silver.py, 547 lignes) "
    "parse le raw_data JSONB collecté en Bronze pour en extraire des colonnes typées. "
    "Idempotence : PK + UNIQUE côté DB + ON CONFLICT … DO UPDATE côté Python."
)

add_h2(doc, '3.1 Inventaire des 5 tables Silver')

add_table(doc,
    headers=['Table', 'Source Bronze', 'Lignes requêtées', 'Dédup', 'Fréquence'],
    rows=[
        ['trafic_boucles_clean', 'trafic_boucles', '200 dernières (>5min)', 'channel_id + measurement_time', '*/5 min'],
        ['tcl_vehicles_clean', 'tcl_vehicles', '200 dernières (sans filtre)', 'line_ref + journey_ref + stop_ref + measurement_time', '*/5 min'],
        ['velov_clean', 'velov', '200 dernières (sans filtre)', 'station_id + measurement_time', '*/5 min'],
        ['meteo_hourly', 'meteo', '200 dernières (sans filtre)', 'measurement_time (PK)', '*/5 min'],
        ['chantiers_actifs', 'chantiers', '200 dernières (sans filtre)', 'chantier_id (⚠ vs fetched_at)', '*/5 min'],
    ],
    col_widths=[Cm(3.5), Cm(2.8), Cm(3.8), Cm(4.5), Cm(2.2)],
)

add_h2(doc, '3.2 Algorithmes de nettoyage')

add_h3(doc, '3.2.1 trafic_boucles_clean (bronze_to_silver.py:79-200)')

add_para(doc, "Pseudo-code de la transformation :")
add_mono(doc, """
  rows = SELECT * FROM bronze.trafic_boucles
         WHERE fetched_at > NOW() - INTERVAL '5 minutes'
         ORDER BY fetched_at DESC LIMIT 200

  for row in rows:
      for feat in row.raw_data["features"]:
          props = feat["properties"]
          if not props.get("code"): continue                # channel_id
          if (channel_id, fetched_at) in seen: continue     # dedup local

          # Parsing regex
          vitesse_kmh = parse_vitesse(props["vitesse"])
              # regex r"^\\s*(\\d+(?:[.,]\\d+)?)\\s*km/h" → float
              # "Vitesse réglementaire" → None

          vitesse_limite_kmh = 50.0                          # ⚠ HARDCODÉ (L132)
          is_sanitary = bool(props.get("est_a_jour", False))

          # Geom : midpoint index du LineString (⚠ cf. D11)
          mid = coords[len(coords)//2]

          INSERT silver.trafic_boucles_clean (...)
          ON CONFLICT (channel_id, measurement_time) DO UPDATE
              SET vitesse_kmh = EXCLUDED.vitesse_kmh, ...
""")

add_h3(doc, '3.2.2 tcl_vehicles_clean (bronze_to_silver.py:293-422)')

add_mono(doc, """
  rows = SELECT * FROM bronze.tcl_vehicles LIMIT 200

  for row in rows:
      delivery = row.raw_data["Siri"]["ServiceDelivery"]
                  ["VehicleMonitoringDelivery"][0]
      for act in delivery.get("VehicleActivity", []):
          mvj = act["MonitoredVehicleJourney"]

          line_ref     = _siri_ref(mvj["LineRef"])              # SIRI 2.0 {"value": "..."} ou string
          if not line_ref: continue

          journey_ref  = (fvj.get("DatedVehicleJourneyRef")
                          or _siri_ref(mvj["VehicleRef"])
                          or "unknown")                          # ⚠ D6
          stop_ref     = _siri_ref(call["StopPointRef"]) or "unknown"

          delay_s      = _parse_siri_delay(mvj["Delay"])
                          # None → 0  ⚠ D3
          direction_ref = _siri_ref(mvj["DirectionRef"])
          lat = loc["Latitude"] if isinstance(loc, dict) else None
          lon = loc["Longitude"] if isinstance(loc, dict) else None

          INSERT silver.tcl_vehicles_clean (...)
          ON CONFLICT (line_ref, journey_ref, stop_ref, measurement_time) DO UPDATE
              SET delay_seconds, fetched_at  (⚠ D15 — ne met pas à jour lat/lon)
""")

add_h3(doc, '3.2.3 velov_clean (bronze_to_silver.py:203-290)')

add_mono(doc, """
  # GBFS unifié Sprint 10+
  if "status" in raw_data and "information" in raw_data:
      info_by_id = {s["station_id"]: s for s in raw_data["information"]}
      stations = ({**st, **info_by_id.get(st["station_id"], {})}
                  for st in raw_data["status"])
  else:
      # Backward-compat legacy
      stations = raw_data.get("data", {}).get("stations", []) \\
                 or raw_data.get("stations", [])

  for st in stations:
      sid = st.get("station_id")
      if not sid: continue
      is_active = bool(st["is_installed"]==1 and
                       st["is_renting"]==1 and
                       st["is_returning"]==1)
      INSERT silver.velov_clean (...)
      ON CONFLICT (station_id, measurement_time) DO UPDATE
""")

add_h3(doc, '3.2.4 meteo_hourly (bronze_to_silver.py:425-477)')

add_mono(doc, """
  # 72 timestamps (1 past + 2 forecast days)
  times = raw_data["hourly"]["time"]
  for i, t in enumerate(times):
      INSERT silver.meteo_hourly
          (measurement_time=t,
           temperature_c=temps[i],                  # ⚠ D2 — colonne inexistante en DDL
           humidity=hums[i], rain_mm=rains[i],      # ⚠ D2
           wind_speed_10m=winds[i], weather_code=codes[i])
      ON CONFLICT (measurement_time) DO UPDATE
  # → 200 × 72 = 14 400 INSERT/UPDATE par run, tous idempotents
""")

add_h3(doc, '3.2.5 chantiers_actifs (bronze_to_silver.py:480-547)')

add_mono(doc, """
  for row in rows:
      for feat in row.raw_data["features"]:
          chantier_id = props.get("id")
          if not chantier_id: continue

          date_debut = props.get("date_debut")
          date_fin   = props.get("date_fin")

          # ⚠ Aucun filtre date_debut ≤ now ≤ date_fin en Python
          # Le filtrage est délégué à un trigger DB (init-db.sql:2445-2457)

          INSERT silver.chantiers_actifs
              (chantier_id, date_debut, date_fin,
               localisation, impact_lines,           # ⚠ D1 — colonnes fantômes
               geom_wgs84, updated_at)              # ⚠ D1
          ON CONFLICT (chantier_id) DO UPDATE       # ⚠ D4 — UNIQUE est (chantier_id, fetched_at)
""")

add_h2(doc, '3.3 Doutes et incohérences — couche Silver')

add_badge(doc, 'CRITIQUE', 'S1 — silver.chantiers_actifs : INSERT impossible avec DDL officiel')
add_para(doc,
    "bronze_to_silver.py:519-540 insère (chantier_id, date_debut, date_fin, localisation, impact_lines, "
    "geom_wgs84, updated_at). Le DDL officiel (init-db.sql:2425-2442) expose (chantier_id, titre, description, "
    "date_debut, date_fin, lat, lon, is_active, raw_data). fetched_at NOT NULL n\'est pas non plus fourni. "
    "L\'INSERT doit lever UndefinedColumn / NotNullViolation à chaque exécution sauf ALTER TABLE non tracké."
)

add_badge(doc, 'CRITIQUE', 'S2 — silver.trafic_boucles_clean.vitesse_kmh vs DDL speed_kmh')
add_para(doc,
    "Code INSERT (bronze_to_silver.py:159) : vitesse_kmh. "
    "DDL (init-db.sql:1282) : speed_kmh. "
    "Code aval (silver_to_gold.py:117) : vitesse_kmh. "
    "Tout le pipeline trafic est probablement cassé en local sauf ALTER appliqué sur VPS."
)

add_badge(doc, 'CRITIQUE', 'S3 — silver.meteo_hourly.temperature_c/rain_mm vs DDL temperature_2m/precipitation')
add_para(doc,
    "Code INSERT (bronze_to_silver.py:453-454) et code aval (db_query.py:773-774, silver_to_gold.py:192-193, "
    "migration 17:92, migration 19:97, migration 21:130) utilisent temperature_c, rain_mm. "
    "DDL expose temperature_2m, precipitation. "
    "Dette documentée 2 fois (CHANGELOG.md:103-106, AUDIT_PIPELINE_2026-06-12.md:182), patch in-place sur VPS "
    "non commité. Le repo n\'est pas la source de vérité sur l\'état réel de la DB."
)

add_badge(doc, 'CRITIQUE', 'S4 — silver.trafic_boucles_clean : DDL sans UNIQUE constraint')
add_para(doc,
    "Code INSERT (L168) déclare ON CONFLICT (channel_id, measurement_time). "
    "DDL officiel (init-db.sql:1279-1288) n\'a aucune contrainte UNIQUE. "
    "ON CONFLICT lève NoUniqueOrPrimaryKeyError ⇒ toutes les INSERT lèvent exception ⇒ aucun trafic en Silver "
    "(sauf UNIQUE ajouté hors-repo)."
)

add_badge(doc, 'MAJEURE', 'S5 — silver.chantiers_actifs.is_active : double DDL')
add_para(doc,
    "init-db.sql:2438 utilise un trigger BEFORE INSERT/UPDATE (idempotent runtime). "
    "migrate_realign_v0.3.1.sql:80-82 utilise une colonne GENERATED ALWAYS AS … STORED — invalide Postgres 16+ "
    "car CURRENT_DATE n\'est pas IMMUTABLE. "
    "Deux DDL différents pour la même table dans le même repo."
)

add_badge(doc, 'MAJEURE', 'S6 — _parse_siri_delay(None) → 0')
add_para(doc,
    "bronze_to_silver.py:295-296 : un retard non-disponible est converti en 0 seconde. "
    "Biais potentiel sur avg_delay_seconds, p90_delay_seconds (silver_to_gold.py:288-322). "
    "À corriger : retourner NULL et propager jusqu\'au Gold (médiane, p90)."
)

add_badge(doc, 'MAJEURE', 'S7 — vitesse_limite_kmh = 50.0 codé en dur')
add_para(doc,
    "bronze_to_silver.py:132 : tous les capteurs (incluant autoroutes A7/A46) reçoivent 50 km/h. "
    "bronze.vitesse_limite_ref existe (init-db.sql:2037) mais aucune jointure. Sprint 10+ TODO non fait."
)

add_badge(doc, 'MAJEURE', 'S8 — Jointure Vélov in-memory (collisions possibles)')
add_para(doc,
    "bronze_to_silver.py:235-238 : info_by_id = {s[\"station_id\"]: s for s in information}. "
    "Si 2 stations ont le même station_id (peu probable mais possible si GBFS renvoie doublons), "
    "le dernier gagne silencieusement."
)

add_badge(doc, 'MAJEURE', 'S9 — journey_ref = "unknown" → collisions sur PK')
add_para(doc,
    "bronze_to_silver.py:377 : journey_ref = (fvj.get(\"DatedVehicleJourneyRef\") or _siri_ref(mvj[\"VehicleRef\"]) "
    "or \"unknown\"). Combiné à stop_ref = \"unknown\" (L381), deux véhicules SIRI différents peuvent "
    "collisionner sur la PK UNIQUE si tous deux ont des refs manquantes et le même line_ref + measurement_time. "
    "Dédup involontaire."
)

add_badge(doc, 'MAJEURE', 'S10 — Filtre temporel Bronze uniquement sur trafic_boucles')
add_para(doc,
    "Lignes 215-220, 343-353, 428-433, 483-488 : ORDER BY fetched_at DESC LIMIT 200 SANS WHERE fetched_at > NOW() - "
    "INTERVAL. Chaque run retraite potentiellement 200 fetches × N features ⇒ ~80-100 min d\'historique."
)

add_badge(doc, 'MINEURE', 'S11 — geom Point mid-idx, pas mid-segment')
add_para(doc,
    "bronze_to_silver.py:147 : mid = coords[len(coords)//2] ⇒ indice du tableau, pas le point au milieu "
    "arithmétique. Si LineString a 3 points [A, B, C], on prend B, pas (A+C)/2."
)

add_badge(doc, 'MINEURE', 'S12 — silver.tcl_vehicles_clean.raw_data jamais alimenté')
add_para(doc,
    "DDL colonne raw_data JSONB (init-db.sql:2405) jamais remplie par l\'INSERT (bronze_to_silver.py:395-399)."
)

add_badge(doc, 'MINEURE', 'S13 — silver.velov_clean.capacity inexistante')
add_para(doc,
    "gold.velov_features.capacity (migrate_realign_v0.3.1.sql:131) référencée, mais silver.velov_clean n\'a "
    "pas cette colonne. Jointure Gold silencieusement NULL."
)

add_badge(doc, 'MINEURE', 'S14 — DAG retries=1 vs règle AGENTS.md')
add_para(doc,
    "AGENTS.md (ligne 47) : DAGs critiques ont retries=0. transform_bronze_to_silver.py:25 : retries=1."
)

add_badge(doc, 'MINEURE', 'S15 — silver.trafic_segments_clean référencée mais inexistante')
add_para(doc,
    "db_query.py:822 SELECT depuis silver.trafic_segments_clean (avec geom_wgs84). "
    "Table jamais créée dans le dump moderne. AUDIT_PIPELINE_2026-06-12.md:411 (DQ-4) le mentionne. "
    "Si la table n\'existe pas, la fonction lève une exception silencieusement attrapée par _df_from_query:130 "
    "qui retourne un DataFrame vide."
)

add_badge(doc, 'MINEURE', 'S16 — migration_021_source_health.sql:130 cible geom_wgs84 sur trafic_boucles_clean')
add_para(doc,
    "La colonne officielle (DDL init-db.sql:1285) est geom, pas geom_wgs84. "
    "La vue gold.v_data_completeness crash si la colonne manque."
)

add_badge(doc, 'MINEURE', 'S17 — Test coverage zéro sur _transform_*')
add_para(doc,
    "tests/integration/test_infrastructure.py ne fait que assert callable(transform_to_silver). "
    "Aucun test unitaire de parsing SIRI, regex Grand Lyon, dedup, idempotence, savepoint."
)

add_badge(doc, 'MINEURE', 'S18 — silver.tcl_vehicles_clean ON CONFLICT incomplet')
add_para(doc,
    "bronze_to_silver.py:400 ON CONFLICT (line_ref, journey_ref, stop_ref, measurement_time) DO UPDATE "
    "SET delay_seconds, fetched_at. Pas de mise à jour de lat, lon, direction_ref, journey_ref, stop_ref lors "
    "d\'un conflict. Si une activité SIRI évolue, seul delay_seconds et fetched_at sont écrasés."
)

page_break(doc)

# =========================================================
# 4. COUCHE GOLD
# =========================================================

add_h1(doc, '4. Couche Gold — Analytique et features')

add_para(doc,
    "La couche Gold contient ~15 tables actives, 4 vues matérialisées rafraîchies en production, 7 vues "
    "non-matérialisées, 1 fonction SQL et 6 tables fantômes (créées, lues par le code mais jamais peuplées). "
    "Le DAG transform_silver_to_gold s\'exécute toutes les 10 minutes (sprint 15+)."
)

add_h2(doc, '4.1 Tables Gold actives (alimentées par DAGs)')

add_table(doc,
    headers=['Table / Vue', 'Type', 'Schedule refresh', 'Consommateurs principaux'],
    rows=[
        ['traffic_features_live', 'TABLE', '*/10 min', 'XGBoost features, widgets Pro_TCL, vues cohérence'],
        ['velov_features', 'TABLE', '*/10 min', 'XGBoost Vélov (dormant)'],
        ['bus_delay_segments', 'TABLE', '*/10 min', 'mv_line_kpis_live, mv_otp_heatmap, mv_bus_traffic_spatial'],
        ['tcl_vehicle_realtime', 'TABLE', '*/10 min', 'mv_multimodal_grid, mv_bus_traffic_spatial, line_kpis'],
        ['infrastructure_bottlenecks', 'TABLE', '*/10 min', 'widgets bottleneck_summary, correlation_matrix'],
        ['trafic_predictions', 'TABLE', '*/15 min', 'widgets trafic (carte), mv_xgb_vs_tomtom, API /predict/traffic'],
        ['xgb_training_set', 'TABLE', 'quotidien 02h30', 'xgboost_speed.train()'],
        ['model_drift_reports', 'TABLE', 'quotidien 05h30', 'widget model_monitoring, health_checks'],
    ],
    col_widths=[Cm(4), Cm(2), Cm(3.5), Cm(8)],
)

add_h2(doc, '4.2 Vues matérialisées stratégiques')

add_table(doc,
    headers=['Vue matérialisée', 'Sprint', 'Schedule', 'Index UNIQUE pour CONCURRENTLY', 'Statut'],
    rows=[
        ['mv_multimodal_grid', '15+ Axe 1', '*/10 min', 'OUI (lat, lon)', '✅ actif'],
        ['mv_bus_traffic_spatial', '15+ Axe 3', '*/10 min', 'OUI (line_ref, hour, lat, lon)', '✅ actif'],
        ['mv_line_kpis_live', '7', 'quotidien 05h', 'OUI (line_ref)', '✅ actif'],
        ['mv_otp_heatmap', '7', 'quotidien 05h', 'OUI (line_id, date, hour)', '✅ actif'],
        ['mv_xgb_vs_tomtom', '15+', '*/30 min', '❌ ABSENT', '⚠ refresh non-CONCURRENTLY (lock 30 min)'],
        ['mv_twgid_to_lyo', '9+', 'aucun DAG', 'OUI (présent)', '🟢 fantôme — créé 1 fois, jamais refresh'],
        ['mv_fact_traffic_pivot', 'init', 'aucun', 'OUI', '🟢 WITH NO DATA — vide'],
    ],
    col_widths=[Cm(4.5), Cm(2.2), Cm(2.8), Cm(4), Cm(3.2)],
)

add_h2(doc, '4.3 Algorithmes des vues phares (Sprint 15+)')

add_h3(doc, '4.3.1 mv_multimodal_grid (migration_017_multimodal_grid.sql:47-139)')

add_para(doc, "Score multimodal 0-10 par cellule géographique (grille 0.01° ≈ 1.1 km) :")
add_mono(doc, """
  score = clamp(0.5 × pct_congestion/10          -- 0..5
              + 0.5 × pct_delayed/10             -- 0..5
              − velov_bonus,                     -- -1.0 si vélos dispo >= 5
              0, 10)

  diagnostic:
    saturated       : pct_congestion > 60 AND pct_delayed > 40
    road_congested  : pct_congestion > 60
    transit_delayed : pct_delayed > 40
    velov_scarce    : bikes_available < 3 AND n_stations > 0
    ok              : reste

  3 CTE FULL OUTER JOIN par (grid_lat, grid_lon) :
    trafic_grid  (gold.traffic_features_live, 1h)
    tcl_grid     (gold.tcl_vehicle_realtime)
    velov_grid   (silver.velov_clean, 15 min)
    meteo        (CROSS JOIN 1 ligne, silver.meteo_hourly)
""")

add_h3(doc, '4.3.2 mv_bus_traffic_spatial (migration_018_bus_traffic_spatial.sql:29-92)')

add_para(doc, "JOIN spatial bus × trafic, fenêtre 7 jours :")
add_mono(doc, """
  bus_positions : GROUP BY (line_ref, HOUR(recorded_at),
                           ROUND(lat, 3), ROUND(lon, 3))
  traffic_zones : GROUP BY (ROUND(lat, 3), ROUND(lon, 3),
                           HOUR(fetched_at))
  -- ⚠ Pas de ST_DWithin, juste arrondi float
  LEFT JOIN sur lat3 = lat3 AND lon3 = lon3 AND hour = hour

  diagnostic (infrastructure_bottlenecks) :
    infra       : bus_delay_sec > 120 AND avg_speed < 25
    operations  : bus_delay_sec > 120 AND (avg_speed >= 25 OR NULL)
    bus_lane_ok : bus_delay_sec <= 120 AND avg_speed < 25
    ok          : reste
""")

add_h3(doc, '4.3.3 fn_network_health_score (migration_019_network_health.sql:30-167)')

add_para(doc, "Score 0-100, redistribution des poids si source indisponible :")
add_mono(doc, """
  Poids initiaux :
    w_traffic = 0.3, w_tcl = 0.3, w_velov = 0.2, w_meteo = 0.2

  Normalisation :
    scale = 1 / (w_traffic + w_tcl + w_velov + w_meteo)
    -- Si meteo indisponible, w_meteo = 0, autres sont renormalisés

  Pénalité météo :
    rain_mm > 5       → 15
    rain_mm > 1       → 8
    temperature_c < 0 → 10
    temperature_c > 35→ 5

  Score final :
    score = clamp(100
                 − pct_congestion × w_traffic × scale
                 − pct_delayed    × w_tcl     × scale
                 − pct_velov_empty× w_velov   × scale
                 − penalty        × w_meteo   × scale,
                 0, 100)

  Diagnostic :
    healthy    : score > 75
    stressed   : score > 50
    degraded   : score > 25
    critical   : sinon

  ⚠ Fenêtres temporelles incohérentes :
    pct_congestion  → 30 min
    pct_delayed     → 30 min
    pct_velov_empty → 15 min
    meteo           → 2h
""")

add_h2(doc, '4.4 Doutes et incohérences — couche Gold')

add_badge(doc, 'CRITIQUE', 'G1 — gold.velov_features : schéma SQL ≠ colonnes INSERTées')
add_para(doc,
    "Code Python (silver_to_gold.py:253-287) INSERT dans colonnes bikes_available, bikes_lag_1, bikes_lag_2, "
    "bikes_lag_3, rolling_mean_3h. "
    "Schéma SQL (migrate_realign_v0.3.1.sql:125-146) : num_bikes_available, lag_30min, lag_60min, rolling_mean_1h. "
    "Conflits : bikes_available vs num_bikes_available (bikes_available n\'existe pas), "
    "bikes_lag_1/2/3 vs lag_30min/60min, rolling_mean_3h vs rolling_mean_1h. "
    "AGENTS.md ligne 24 le mentionne comme dette technique connue non résolue."
)

add_badge(doc, 'CRITIQUE', 'G2 — gold.velov_predictions lue par db_query mais jamais peuplée')
add_para(doc,
    "Table définie (migrate_realign_v0.3.1.sql:148-160). "
    "Aucun INSERT trouvé dans le repo (grep 0 hit). "
    "db_query.get_velov_predictions:421 la SELECT. "
    "Conséquence : widget cassé silencieusement via _df_from_query:130 → DataFrame vide."
)

add_badge(doc, 'CRITIQUE', 'G3 — gold.amenagements_history, gold.fact_correlation_matrix, gold.mv_kpis_12_months lues mais inexistantes')
add_para(doc,
    "db_query.get_amenagements_passes:888, get_correlation_matrix:833, get_kpis_12_months:867 font des SELECT "
    "sur des tables qui n\'existent pas dans le dump moderne. "
    "Aucune trace d\'ALTER TABLE ou CREATE TABLE pour ces tables. "
    "Soit à DROP les appels, soit à créer les tables."
)

add_badge(doc, 'CRITIQUE', 'G4 — gold.channels_ref JOIN dans v_coherence_tomtom_vs_grandlyon non peuplée')
add_para(doc,
    "migration_14_gold_coherence_tomtom_v2.sql:69 fait FROM gold.channels_ref cr. "
    "Aucun INSERT actif dans le code. "
    "Si la table est vide, la vue ne renvoie aucune paire → v_tomtom_gl_drift vide → widget tomtom_coherence "
    "(Pro_TCL) affiche 0 lignes. AGENTS.md dette schéma non résolue."
)

add_badge(doc, 'CRITIQUE', 'G5 — Trigger lat/lon jamais déclenché (dim_spatial_grid_mapping vide)')
add_para(doc,
    "audit_dim_spatial_writers.sql:22-27 trigger trg_dim_spatial_has_lat_lon refuse INSERT/UPDATE si "
    "properties_twgid !~ '^[0-9]+$' ET lat/lon NULL. "
    "Mais build_spatial_mapping.py:60 fait continue si lat/lon NULL → skip silencieux → la table reste vide. "
    "Le trigger ne se déclenche jamais pour ces rows. AGENTS.md dette non résolue en Sprint 16."
)

add_badge(doc, 'MAJEURE', 'G6 — _BOTTLENECK_SQL et migration_018 sémantiques différentes')
add_para(doc,
    "_BOTTLENECK_SQL (silver_to_gold.py:398-443) : agrège par HEURE GLOBALE (avg_speed de TOUT Lyon). "
    "migration_018_bus_traffic_spatial.sql : JOIN spatialisé (avg_speed par zone). "
    "Les 2 sources coexistent avec des sémantiques différentes. Les widgets correlation_matrix et "
    "load_infra_bottlenecks continuent de lire l\'ancienne table. Pas de bascule vers Option A (remplacement) "
    "prévue avant 7 jours de données MV (commentaire migration_018:19-20)."
)

add_badge(doc, 'MAJEURE', 'G7 — mv_xgb_vs_tomtom pas d\'index UNIQUE')
add_para(doc,
    "migration_020_xgb_vs_tomtom.sql:99-104 : 3 index B-tree mais aucun UNIQUE. "
    "Conséquence documentée (refresh_xgb_vs_tomtom.py:53-56) : le refresh ne peut PAS être CONCURRENTLY ⇒ "
    "lock exclusif toutes les 30 min. Le widget backtest_dashboard (Pro_7) unavailable 5-10s toutes les 30 min."
)

add_badge(doc, 'MAJEURE', 'G8 — Fenêtres temporelles incohérentes dans fn_network_health_score')
add_para(doc,
    "pct_congestion 30 min, pct_delayed 30 min, pct_velov_empty 15 min, meteo 2h. "
    "Une panne TCL de 20 min sera invisible ; un pic meteo vieux de 90 min continuera à pénaliser. "
    "À documenter ou unifier à 30 min."
)

add_badge(doc, 'MINEURE', 'G9 — gold.mv_twgid_to_lyo jamais rafraîchie')
add_para(doc,
    "create_mv_twgid_to_lyo.sql:22-24 DROP+CREATE mais aucun DAG ne fait REFRESH. "
    "Aucun consommateur trouvé (grep mv_twgid_to_lyo ⇒ 0 hit hors définition). "
    "Code mort. À DROP ou à refresh par un DAG."
)

add_badge(doc, 'MINEURE', 'G10 — gold.trafic_predictions.x_2154, y_2154 toujours NULL')
add_para(doc,
    "Schéma (init-db.sql:1221-1222) mais dag_inference_xgboost.py:162-169 ne les insère pas. "
    "À DROP les colonnes ou à les peupler."
)

add_badge(doc, 'MINEURE', 'G11 — gold.bus_delay_segments.weather_code toujours NULL')
add_para(doc,
    "silver_to_gold.py:305 insère NULL::int pour weather_code. À DROP la colonne ou intégrer météo dans _BUS_DELAY_SQL."
)

add_badge(doc, 'MINEURE', 'G12 — gold.multimodal_status_grid fantôme sur bronze.pvotrafic_snapshots')
add_para(doc,
    "init-db.sql:943-998 crée cette vue sur une table qui n\'est plus dans le dump moderne. "
    "Aucun consommateur. À DROP ou remplacer par mv_multimodal_grid."
)

add_badge(doc, 'MINEURE', 'G13 — v_recent_alerts citée dans db_query.py:791 mais inexistante')
add_para(doc,
    "Le commentaire dit « mock — sera remplacé par gold.v_recent_alerts ». La vue n\'existe pas. "
    "Le code lit silver.chantiers_actifs. Commentaire contradictoire avec la politique zéro mock Sprint 8."
)

add_badge(doc, 'MINEURE', 'G14 — silver.trafic_segments_clean lue par get_segments:822 mais inexistante')
add_para(doc,
    "cf. S15. À DROP ou créer."
)

add_badge(doc, 'MINEURE', 'G15 — Pas d\'index (channel_id, computed_at DESC) sur gold.traffic_features_live')
add_para(doc,
    "db_query.get_traffic_for_node:200 filtre WHERE computed_at >= NOW() - INTERVAL '2 hours' ⇒ seq scan."
)

add_badge(doc, 'MINEURE', 'G16 — gold.tarifs_modes non lue par le code actif')
add_para(doc,
    "Migration 016 seed 16 produits. src/routing/eco_calculator.py reste 100% en constantes hardcodées. "
    "Le commentaire migration_016:17-22 l\'avoue (« Phase 1 : constantes … Phase 2+ : DB »). "
    "La migration a devancé le code."
)

add_badge(doc, 'MINEURE', 'G17 — Jointures FULL OUTER ambiguës dans mv_multimodal_grid')
add_para(doc,
    "3 FULL OUTER JOIN sur (grid_lat, grid_lon). Pas de garde EXISTS. "
    "Risque mineur de dédoublement quand silver.velov_clean arrondit différemment de gold.traffic_features_live."
)

add_badge(doc, 'COSMETIQUE', 'G18 — 11 tables Gold legacy fantômes')
add_para(doc,
    "h3_trafic_live, h3_trafic_predictions, features_traffic, road_importance_ref, sensor_road_importance, "
    "channel_tomtom_mapping, stgcn_predictions_live, dim_temps, etc. "
    "Aucun consommateur. À DROP en Sprint 17+."
)

add_badge(doc, 'COSMETIQUE', 'G19 — gold.infrastructure_bottlenecks.lat/lon synthétiques via HASHTEXT')
add_para(doc,
    "silver_to_gold.py:438-439 insère lat/lon calculés via HASHTEXT(line_ref), pas géographique. "
    "data_loader._approx_lonlat_from_channel_id:71-93 le reconnaît explicitement. À calculer via lookup réel."
)

page_break(doc)

# =========================================================
# 5. MODÈLES ML
# =========================================================

add_h1(doc, '5. Modèles ML')

add_para(doc,
    "Trois modèles sont déclarés dans le projet : XGBoost Speed (production), XGBoost Vélov (dormant) et "
    "ST-GRU-GNN (challenger, désactivé). Seul XGBoost Speed tourne effectivement et écrit dans gold.trafic_predictions."
)

add_h2(doc, '5.1 Tableau récapitulatif des 3 modèles')

add_table(doc,
    headers=['Modèle', 'Statut', 'Horizon(s)', 'Retrain', 'Inférence', 'MLflow'],
    rows=[
        ['XGBoost Speed (H+1h)', '✅ Production', '60 min', 'quotidien 03h00', '*/15 min', 'exp xgboost_speed, model 1.2.0'],
        ['XGBoost Vélov (H+30/H+60)', '⚠ Dormant (features cassées)', '30 + 60 min', 'hourly :50 (DAG présent, cassé)', 'aucune (DAG jamais créé)', 'exp xgboost_velov'],
        ['ST-GRU-GNN (SpatioTemporalGCN)', '❌ Inactif (DAG paused)', '5/15/30/60/180/360 min', 'quotidien 03h00 (paused)', 'aucune', 'exp stgcn_traffic, model 0.3.0'],
    ],
    col_widths=[Cm(4.5), Cm(3), Cm(2.5), Cm(3), Cm(2.5), Cm(3.5)],
)

add_h2(doc, '5.2 XGBoost Speed — features (source de vérité)')

add_para(doc, "11 features déclarées dans xgboost_speed.py:50-62 :")
add_table(doc,
    headers=['#', 'Feature', 'Type', 'Source (gold.traffic_features_live.col)', 'Présent'],
    rows=[
        ['1', 'speed_kmh', 'float', 'speed_kmh (init-db.sql:1010)', '✅'],
        ['2', 'lag_1', 'float', 'lag_1 (init-db.sql:1012)', '✅'],
        ['3', 'lag_2', 'float', 'lag_2 (init-db.sql:1013)', '✅'],
        ['4', 'lag_3', 'float', 'lag_3 (init-db.sql:1014)', '✅'],
        ['5', 'rolling_mean_3', 'float', 'rolling_mean_3 (init-db.sql:1017)', '✅'],
        ['6', 'sin_hour', 'float', 'sin_hour (init-db.sql:1021)', '✅'],
        ['7', 'cos_hour', 'float', 'cos_hour (init-db.sql:1022)', '✅'],
        ['8', 'temperature_2m', 'float', 'temperature_2m (init-db.sql:1026)', '✅'],
        ['9', 'precipitation', 'float', 'precipitation (init-db.sql:1027)', '✅'],
        ['10', 'is_vacances', 'bool', 'is_vacances (init-db.sql:1038)', '✅'],
        ['11', 'is_ferie', 'bool', 'is_ferie (init-db.sql:1039)', '✅'],
    ],
    col_widths=[Cm(0.8), Cm(3.2), Cm(2), Cm(7), Cm(2)],
)
add_para(doc, "✅ Cohérent avec le schéma SQL. gold.xgb_training_set (create_xgb_training_set.sql:38-64) a "
              "exactement les mêmes 11 colonnes.", italic=True)

add_h2(doc, '5.3 XGBoost Vélov — features (casse-tête)')

add_para(doc, "11 features déclarées dans xgboost_velov.py:25-37 :")
add_table(doc,
    headers=['#', 'Feature (code)', 'Présent dans gold.velov_features ?', 'Colonne réelle en DB'],
    rows=[
        ['1', 'station_id_encoded', '✅', 'station_id_encoded'],
        ['2', 'bikes_lag_1', '❌', 'lag_30min (migrate_realign_v0.3.1.sql:141)'],
        ['3', 'bikes_lag_2', '❌', 'lag_60min (migrate_realign_v0.3.1.sql:142)'],
        ['4', 'bikes_lag_3', '❌', '∅'],
        ['5', 'rolling_mean_3h', '❌', 'rolling_mean_1h (migrate_realign_v0.3.1.sql:143)'],
        ['6', 'hour_sin', '✅', 'hour_sin'],
        ['7', 'hour_cos', '✅', 'hour_cos'],
        ['8', 'temperature_c', '✅', 'temperature_c'],
        ['9', 'rain_mm', '✅', 'rain_mm'],
        ['10', 'is_vacances', '✅', 'is_vacances'],
        ['11', 'is_ferie', '✅', 'is_ferie'],
    ],
    col_widths=[Cm(0.8), Cm(3.5), Cm(4), Cm(6.5)],
)
add_para(doc, "Tout train_one(30) ou train_one(60) plantera sur KeyError si ce DAG est activé. "
              "Dette technique connue documentée mais non résolue (AGENTS.md ligne 24).", italic=True, color=C_CRITIQUE)

add_h2(doc, '5.4 ST-GRU-GNN — architecture')

add_mono(doc, """
  input  : (seq_len, 5 canaux)  -- speed_norm, hour_sin, hour_cos, day_sin, day_cos
  hidden : 128
  archi  : GRU(1 layer) → 2× GCNConv(128→128) + LeakyReLU(0.2)
           + LayerNorm + skip connections → Linear(128→1)
  loss   : MSELoss
  optim  : Adam lr=1e-3 + ReduceLROnPlateau (patience=3, factor=0.5)
  epochs : 50, batch=16, early_stopping=5
  quality gate : new_mae ≤ prev_mae × 1.15
                  (LYONFLOW_STRICT_QUALITY env var)

  Graphe : H3 res 13, K=2 grid_disk
           ~1520 nœuds × ~9540 arêtes
  Source : gold.fact_traffic_series (5 canaux)
           -- alimentée par dags/legacy_github/dag_pipeline.py:613 (DAG désactivé)
""")

add_h2(doc, '5.5 Doutes et incohérences — Modèles ML')

add_badge(doc, 'CRITIQUE', 'M1 — XGBoost Vélov : features ne matchent pas la table')
add_para(doc,
    "Le modèle lit bikes_lag_1/2/3, rolling_mean_3h (xgboost_velov.py:27-30, 173-174) ; "
    "la table gold.velov_features a lag_30min, lag_60min, rolling_mean_1h. "
    "Tout train_one plantera sur KeyError si ce DAG tourne. AGENTS.md § « Dette technique connue » le mentionne."
)

add_badge(doc, 'CRITIQUE', 'M2 — DAGs XGBoost legacy = doublons actifs')
add_para(doc,
    "dags/ml/retrain_xgboost.py:101-115 crée un DAG retrain_xgboost_speed schedulé 25 * * * * (hourly) qui "
    "appelle model.train_one(5), train_one(60), train_one(180), train_one(360). "
    "Mais le code XGBoost force horizon_minutes=60 uniquement. "
    "À chaque heure :25, le DAG s\'exécute, lance 4 train_one qui retournent des fallbacks et… ré-entraînent "
    "4 modèles .pkl obsolètes sur disque. Pollution."
)

add_badge(doc, 'CRITIQUE', 'M3 — MLflow : pas d\'URI effective')
add_para(doc,
    "mlflow_integration.py:98, 105 initialise avec tracking_uri=os.getenv(\"MLFLOW_TRACKING_URI\", "
    "\"http://localhost:5000\"). Si l\'env var est vide, MLflow écrit dans ./mlruns (filesystem local) — "
    "les runs ne vont PAS dans le serveur MLflow central. "
    "Le test à xgboost_speed.py:193 est if os.getenv(\"MLFLOW_TRACKING_URI\", \"\") != \"\" — tracking "
    "désactivé par défaut. Conséquence : load() (xgboost_speed.py:106-115) tombe systématiquement en "
    "fallback disque."
)

add_badge(doc, 'CRITIQUE', 'M4 — STGCNWrapper jamais branché en prod')
add_para(doc,
    "Aucun DAG Airflow n\'appelle STGCNWrapper.predict(). "
    "Le handle STGCNModelHandle (model_registry.py:200-244) est mort-né : "
    "LYONFLOW_MODELS_ACTIVE=both (défaut) → get_active_model() retourne xgboost (champion), "
    "get_challenger_model() retourne stgcn mais personne ne le consomme. "
    "Vérifié : aucune occurrence de STGCNWrapper ou STGCNModelHandle dans src/api/, dashboard/ ou DAGs."
)

add_badge(doc, 'MAJEURE', 'M5 — Fallback silencieux 30 km/h')
add_para(doc,
    "xgboost_speed.predict():286-320 retourne des fallbacks hardcodés (30 km/h, model_version=\"0.0.0\") "
    "si horizon != 60 OU si modèle non chargé OU si features vides. "
    "Aucun widget ne distingue vraie prédiction vs fallback 30.0. Le widget model_monitoring affiche "
    "« XGBoost Speed H+1h » sans check de version. Viole la règle « fail loud » du Sprint 8."
)

add_badge(doc, 'MAJEURE', 'M6 — model_version mismatch (1.0.0 vs 1.2.0)')
add_para(doc,
    "xgboost_speed.py:208 log model_version=1.2.0 dans MLflow lors du train_one. "
    "Mais :334 retourne model_version=1.0.0 lors du predict. "
    "Deux sources de vérité contradictoires. Le widget model_monitoring:394 affiche via "
    "get_latest_run(...).params.model_version (MLflow) = 1.2.0, mais le model_version dans "
    "gold.trafic_predictions (dag_inference_xgboost.py:151) sera 1.0.0."
)

add_badge(doc, 'MAJEURE', 'M7 — Mauvais model_dir lookup (chemin mort dans ModelRegistry)')
add_para(doc,
    "src/ml/model_registry.py:172 teste f\"xgb_speed_h{self.horizon_min}.json\" mais le code écrit en .pkl "
    "(xgboost_speed.py:189). Le is_available() retournera toujours False — le singleton ModelRegistry ne "
    "charge jamais XGBoost sauf si on contourne via model.load() directement."
)

add_badge(doc, 'MAJEURE', 'M8 — MLflow register_model cassé en transition')
add_para(doc,
    "mlflow_integration.py:255-260 : client.create_model_version(name=model_name, source=uri, run_id=run_id). "
    "source doit être un artifact URI valide du run, mais le uri construit est f\"runs:/{run_id}/"
    "{model_name}.pkl\" (L251) — ce fichier n\'existe pas comme artifact du run. "
    "MLflow va créer une ModelVersion pointant vers un artifact inexistant."
)

add_badge(doc, 'MINEURE', 'M9 — daily_drift_report vs build_xgb_training_set : 2 stratégies différentes')
add_para(doc,
    "daily_drift_report.py utilise run_drift_report() (drift_detector.py:37 référence xgb_speed_kmh). "
    "Mais build_xgb_training_set.py:191 calcule du PSI sur target_speed, speed_kmh, temperature_2m, "
    "precipitation (PAS tomtom_speed_kmh). "
    "Deux stratégies de drift différentes, aucune ne couvre les mêmes features."
)

add_badge(doc, 'MINEURE', 'M10 — 2 writers sur gold.model_drift_reports')
add_para(doc,
    "build_xgb_training_set.py:222 et daily_drift_report.py:49 font INSERT. "
    "Risque de collision INSERT sans ON CONFLICT."
)

add_badge(doc, 'MINEURE', 'M11 — STGCNTrainer 6 horizons × VPS 12 Go RAM')
add_para(doc,
    "STGCNTrainer.DEFAULT_HORIZONS_MIN = (5, 15, 30, 60, 180, 360). "
    "Sur VPS (12 Go RAM), batch_size=16, 6 horizons = 6× chargement + 6× forward pass. "
    "Estimé 30-60 min par train_all() sur CPU. Aucune trace de mesure de perf."
)

add_badge(doc, 'MINEURE', 'M12 — Pas de model_card pour Vélov et ST-GCN')
add_para(doc,
    "Seul XGBoost Speed a model_card.py (xgboost_speed.py:222-247). Inhérent à M1, M4."
)

page_break(doc)

# =========================================================
# 6. ROUTING MULTIMODAL
# =========================================================

add_h1(doc, '6. Routing multimodal')

add_para(doc,
    "Le module de routing (src/routing/) implémente une recommandation multimodale pour 3 modes : voiture, "
    "transports en commun (TC), vélo (Vélov). Il s\'appuie sur le pathfinder Dijkstra/A* sur graphe routier H3 "
    "et l\'éco-calculateur pour les facteurs CO2 et tarifs."
)

add_h2(doc, '6.1 Algorithme de voiture — Dijkstra / A*')

add_mono(doc, """
  # src/routing/pathfinder.py:99 — nx.astar_path (PAS Dijkstra, le nom est trompeur)
  # Heuristique : haversine bornée à 50 km/h
  # Graphe :
  #   gold.dim_spatial_grid_mapping (1520 nœuds H3 res 13)
  #   gold.dim_gnn_adjacency (K=2, ~4072 arêtes)
  # Poids :
  #   travel_time_s = length_m / (speed_kmh × 1000/3600)
  #   speed = min(get_node_speed(u), get_node_speed(v))
  # ⚠ get_node_speed ignore horizon_minutes (graph.py:290-298)
  #   → routage "traffic-aware" est en fait H+0 toujours
  # ⚠ _compute_confidence retourne 0.85 hardcodé (pathfinder.py:172-174)
  # ⚠ build_spatial_mapping.py:60 skip silencieux si lat/lon NULL
""")

add_h2(doc, '6.2 Algorithme Vélov — smart routing')

add_para(doc, "Implémenté dans pathfinder_multimodal.py:254-537 :")
add_mono(doc, """
  1. Snap GPS → lieu_id via referentiel.lieux_lyon + haversine (L304-316)
  2. Cherche top 3 bornes scorées via referentiel.v_lieux_velov_smart (Sprint VPS-6)
  3. _pick_best (L333-345) : priorité OK > FAIBLE > 1ère dispo
  4. 3 segments : marche → vélo (haversine ou Dijkstra) → marche
  5. Diagnostics si VIDE ou PLEINE (L405-449)
""")

add_h2(doc, '6.3 Algorithme TC')

add_para(doc, "Implémenté dans pathfinder_multimodal.py:928-1105 :")
add_mono(doc, """
  1. lignes(O) ∩ lignes(D) non vide → trajet direct (L998-1022)
     Meilleure ligne = min(rank_O + rank_D)
  2. Sinon → 1 hub (parmi 21 lieux actifs) minimisant durée (L1024-1094)
  3. Cadence : referentiel.lieux_calendrier (L856-882)
  4. Retard : gold.bus_delay_segments moyen 7j (L884-901)
  5. Durée : walk_to + cadence/2 + drive + delay_avg + walk_from
  6. Pénalité correspondance : +3 min (L1075)

  ⚠ Vitesse TC hardcodée :
    metro=35, tram=20, bus=15, funicular=18 km/h (L691-696)
    Source : commentaire uniquement, pas de table DB
  ⚠ Détection is_vacation simplifiée (L779) :
    1/7→31/8 + 15/12→5/1 — ne matche pas la zone A/B/C officielle
""")

add_h2(doc, '6.4 Éco-calculateur — facteurs CO2 et tarifs')

add_table(doc,
    headers=['Constante', 'Valeur', 'Source', 'Fichier:ligne'],
    rows=[
        ['VOITURE_CO2_G_PER_KM', '193 g/km', 'ADEME Base Carbone 2024 VP essence', 'eco_calculator.py:51'],
        ['VOITURE_CONGESTION_PENALTY', '× 1.4', 'ADEME étude impact', 'eco_calculator.py:54'],
        ['TCL_CO2_G_PER_KM', '35 g/km', 'SYTRAL/ADEME mix métro-tram-bus', 'eco_calculator.py:57'],
        ['VELOV_CO2_G_PER_KM', '0', 'scope opérationnel', 'eco_calculator.py:61'],
        ['VOITURE_FUEL_PRICE_EUR', '1.85 €/L', '—', 'eco_calculator.py:53'],
        ['TCL_TICKET_UNITAIRE_EUR', '2.05 €', '—', 'eco_calculator.py:58'],
        ['VELOV_COST_EUR', '0 (abonné < 30 min)', '—', 'eco_calculator.py:62'],
        ['CALORIES_PER_KM (velov)', '46', 'MET ADEME/INSERM', 'eco_calculator.py:66'],
        ['_CONGESTION_SPEED_THRESHOLD_KMH', '25.0', '—', 'eco_calculator.py:69'],
        ['_TIME_VALUE_EUR_PER_MIN', '0.30', 'CEREMA 2023', 'eco_calculator.py:73'],
    ],
    col_widths=[Cm(5), Cm(3.5), Cm(5), Cm(4)],
)

add_h2(doc, '6.5 Doutes et incohérences — Routing')

add_badge(doc, 'MAJEURE', 'R1 — recommend_mode implémente un scoring différent de la spec')
add_para(doc,
    "Contre toute attente, le code eco_calculator.py:225-281 n\'utilise PAS de formule pondérée 50/30/20 "
    "(50% temps + 30% coût + 20% CO2). "
    "Le critère \"temps\" ignore le coût, le critère \"cout\" ignore le CO2. "
    "Le CO2 n\'est jamais intégré au score. Le critere n\'accepte que \"temps\" ou \"cout\" (ValueError sinon). "
    "recommend_mode n\'est pas utilisé en production (mode_comparison.py:307 confirme)."
)

add_badge(doc, 'MAJEURE', 'R2 — Routage "traffic-aware H+1h" ignore l\'horizon')
add_para(doc,
    "Le routage \"traffic-aware H+1h\" annoncé partout (pathfinder_multimodal.py:552, 567) n\'utilise JAMAIS "
    "l\'horizon. Le commentaire graph.py:297 le reconnaît : « # Sprint 6+ : intégrer gold.trafic_predictions "
    "pour horizon > 0 ». Si un user demande horizon_minutes=60 à plan_car_trip, l\'API retourne des vitesses H+0, "
    "pas H+1h. Incohérence avec la spec « Routage H+1h »."
)

add_badge(doc, 'MAJEURE', 'R3 — Mauvais model_dir lookup (chemin mort dans ModelRegistry)')
add_para(doc,
    "src/ml/model_registry.py:172 teste f\"xgb_speed_h{self.horizon_min}.json\" mais le code écrit en .pkl. "
    "Le is_available() retournera toujours False."
)

add_badge(doc, 'MINEURE', 'R4 — _compute_confidence = 0.85 hardcodé')
add_para(doc,
    "pathfinder.py:172-174 : _compute_confidence retourne 0.85 hardcodé. "
    "Commentaire « Sprint 6+ : query DB pour vérifier fraîcheur réelle », jamais fait."
)

add_badge(doc, 'MINEURE', 'R5 — channel_id retourné = f\"{u}→{v}\"')
add_para(doc,
    "pathfinder.py:127 : le channel_id retourné par les segments est f\"{u}→{v}\" (concaténation), "
    "pas un vrai LYO00xxx — perte d\'information."
)

add_badge(doc, 'MINEURE', 'R6 — gold.tarifs_modes jamais lue par le code actif')
add_para(doc,
    "cf. G16. La table est seedée (migration_016) mais le code reste en constantes hardcodées. "
    "La migration a devancé le code."
)

add_badge(doc, 'MINEURE', 'R7 — TTL cache graph = 300 s hardcodé')
add_para(doc,
    "graph.py:45 : TTL cache 300s hardcodé, pas configurable par env."
)

add_badge(doc, 'MINEURE', 'R8 — Backend pathfinder rebuild du graphe incomplet')
add_para(doc,
    "graph.py:193-198 : « Keep only the largest connected component so get_nearest_node never returns an "
    "isolated node » — supprime les nœuds isolés APRÈS le calcul, donc un user dont la destination est sur un "
    "nœud isolé aura get_nearest_node qui retourne un nœud éloigné → itinéraire aberrant."
)

add_badge(doc, 'MINEURE', 'R9 — plan_transit_trip charge tous les lieux (acceptable jusqu\'à 21)')
add_para(doc,
    "pathfinder_multimodal.py:1025 : SELECT lieu_id, name FROM referentiel.lieux_lyon WHERE is_active = TRUE. "
    "21 lieux donc 21 boucles × queries cadences/retards. Acceptable en l\'état mais inacceptable si on scale à 200 lieux."
)

add_badge(doc, 'MINEURE', 'R10 — _get_current_day_type_and_bucket sans timezone')
add_para(doc,
    "pathfinder_multimodal.py:797-804 : lit datetime.now() sans timezone → risque d\'incohérence si l\'env "
    "Docker n\'est pas en Europe/Paris. Sprint 14+ n\'a pas forcé TZ=Europe/Paris dans le container."
)

page_break(doc)

# =========================================================
# 7. SYNTHÈSE DES DOUTES
# =========================================================

add_h1(doc, '7. Synthèse des doutes identifiés')

add_para(doc, "Tableau récapitulatif trié par sévérité :")

add_table(doc,
    headers=['#', 'ID', 'Couche', 'Sévérité', 'Sujet'],
    rows=[
        ['1', 'D1', 'Bronze', 'CRITIQUE', 'CalendrierScolaire + JoursFeries INSERT cassé (colonnes fetched_at/raw_data absentes)'],
        ['2', 'D7', 'Bronze', 'CRITIQUE', 'run() swallow exceptions → tâches Airflow "vertes" en erreur réelle'],
        ['3', 'D15', 'Bronze', 'CRITIQUE', 'gold.v_source_health référence fetched_at sur tables inexistantes'],
        ['4', 'S1', 'Silver', 'CRITIQUE', 'silver.chantiers_actifs INSERT impossible avec DDL officiel (colonnes fantômes)'],
        ['5', 'S2', 'Silver', 'CRITIQUE', 'silver.trafic_boucles_clean.vitesse_kmh vs DDL speed_kmh'],
        ['6', 'S3', 'Silver', 'CRITIQUE', 'silver.meteo_hourly.temperature_c/rain_mm vs DDL temperature_2m/precipitation'],
        ['7', 'S4', 'Silver', 'CRITIQUE', 'silver.trafic_boucles_clean DDL sans UNIQUE constraint'],
        ['8', 'G1', 'Gold', 'CRITIQUE', 'gold.velov_features schéma SQL ≠ colonnes INSERTées'],
        ['9', 'G2', 'Gold', 'CRITIQUE', 'gold.velov_predictions lue par db_query mais jamais peuplée'],
        ['10', 'G3', 'Gold', 'CRITIQUE', 'gold.amenagements_history, gold.fact_correlation_matrix, gold.mv_kpis_12_months lues mais inexistantes'],
        ['11', 'G4', 'Gold', 'CRITIQUE', 'gold.channels_ref JOIN dans v_coherence_tomtom_vs_grandlyon non peuplée'],
        ['12', 'G5', 'Gold', 'CRITIQUE', 'Trigger lat/lon jamais déclenché (dim_spatial_grid_mapping vide)'],
        ['13', 'M1', 'ML', 'CRITIQUE', 'XGBoost Vélov features ne matchent pas la table (KeyError garanti)'],
        ['14', 'M2', 'ML', 'CRITIQUE', 'DAGs XGBoost legacy = doublons actifs (retrain_xgboost hourly)'],
        ['15', 'M3', 'ML', 'CRITIQUE', 'MLflow : pas d\'URI effective en prod'],
        ['16', 'M4', 'ML', 'CRITIQUE', 'STGCNWrapper jamais branché en prod'],
        ['17', 'D2', 'Bronze', 'MAJEURE', 'bronze.trafic_vitesse_brute DROP mentionnée mais pas exécutée'],
        ['18', 'D8', 'Bronze', 'MAJEURE', 'Docstrings mensongères sur les fréquences'],
        ['19', 'D9', 'Bronze', 'MAJEURE', 'chantiers pollue Bronze à 100k rows/jour'],
        ['20', 'D11', 'Bronze', 'MAJEURE', '_validate_table whitelist inclut calendriers_scolaire + jours_feries'],
        ['21', 'S5', 'Silver', 'MAJEURE', 'silver.chantiers_actifs.is_active : double DDL'],
        ['22', 'S6', 'Silver', 'MAJEURE', '_parse_siri_delay(None) → 0'],
        ['23', 'S7', 'Silver', 'MAJEURE', 'vitesse_limite_kmh = 50.0 codé en dur'],
        ['24', 'S8', 'Silver', 'MAJEURE', 'Jointure Vélov in-memory (collisions possibles)'],
        ['25', 'S9', 'Silver', 'MAJEURE', 'journey_ref = "unknown" → collisions sur PK'],
        ['26', 'S10', 'Silver', 'MAJEURE', 'Filtre temporel Bronze uniquement sur trafic_boucles'],
        ['27', 'G6', 'Gold', 'MAJEURE', '_BOTTLENECK_SQL et migration_018 sémantiques différentes'],
        ['28', 'G7', 'Gold', 'MAJEURE', 'mv_xgb_vs_tomtom pas d\'index UNIQUE'],
        ['29', 'G8', 'Gold', 'MAJEURE', 'Fenêtres temporelles incohérentes dans fn_network_health_score'],
        ['30', 'M5', 'ML', 'MAJEURE', 'Fallback silencieux 30 km/h'],
        ['31', 'M6', 'ML', 'MAJEURE', 'model_version mismatch (1.0.0 vs 1.2.0)'],
        ['32', 'M7', 'ML', 'MAJEURE', 'Mauvais model_dir lookup (chemin mort)'],
        ['33', 'M8', 'ML', 'MAJEURE', 'MLflow register_model cassé en transition'],
        ['34', 'R1', 'Routing', 'MAJEURE', 'recommend_mode implémente un scoring différent de la spec'],
        ['35', 'R2', 'Routing', 'MAJEURE', 'Routage "traffic-aware H+1h" ignore l\'horizon'],
        ['36', 'D3', 'Bronze', 'MINEURE', 'Colonnes extracted non peuplées (sauf TomTom)'],
        ['37', 'D4', 'Bronze', 'MINEURE', 'bronze.jours_feries sans index ni PK'],
        ['38', 'D5', 'Bronze', 'MINEURE', 'Doublon bronze.tomtom_flow vs bronze.tomtom_traffic'],
        ['39', 'D6', 'Bronze', 'MINEURE', '6 tables Bronze orphelines'],
        ['40', 'D10', 'Bronze', 'MINEURE', 'Pas de métriques Prometheus sur les collecteurs'],
        ['41', 'D14', 'Bronze', 'MINEURE', 'TomTomTrafficFlow doublement planifié'],
        ['42', 'S11', 'Silver', 'MINEURE', 'geom Point mid-idx, pas mid-segment'],
        ['43', 'S12', 'Silver', 'MINEURE', 'silver.tcl_vehicles_clean.raw_data jamais alimenté'],
        ['44', 'S13', 'Silver', 'MINEURE', 'silver.velov_clean.capacity inexistante'],
        ['45', 'S14', 'Silver', 'MINEURE', 'DAG retries=1 vs règle AGENTS.md'],
        ['46', 'S15', 'Silver', 'MINEURE', 'silver.trafic_segments_clean référencée mais inexistante'],
        ['47', 'S16', 'Silver', 'MINEURE', 'migration_021_source_health.sql:130 cible geom_wgs84'],
        ['48', 'S17', 'Silver', 'MINEURE', 'Test coverage zéro sur _transform_*'],
        ['49', 'S18', 'Silver', 'MINEURE', 'silver.tcl_vehicles_clean ON CONFLICT incomplet'],
        ['50', 'G9', 'Gold', 'MINEURE', 'gold.mv_twgid_to_lyo jamais rafraîchie'],
        ['51', 'G10', 'Gold', 'MINEURE', 'gold.trafic_predictions.x_2154, y_2154 toujours NULL'],
        ['52', 'G11', 'Gold', 'MINEURE', 'gold.bus_delay_segments.weather_code toujours NULL'],
        ['53', 'G12', 'Gold', 'MINEURE', 'gold.multimodal_status_grid fantôme'],
        ['54', 'G13', 'Gold', 'MINEURE', 'v_recent_alerts citée mais inexistante'],
        ['55', 'G14', 'Gold', 'MINEURE', 'silver.trafic_segments_clean lue mais inexistante'],
        ['56', 'G15', 'Gold', 'MINEURE', 'Pas d\'index (channel_id, computed_at DESC) sur gold.traffic_features_live'],
        ['57', 'G16', 'Gold', 'MINEURE', 'gold.tarifs_modes non lue par le code actif'],
        ['58', 'G17', 'Gold', 'MINEURE', 'Jointures FULL OUTER ambiguës dans mv_multimodal_grid'],
        ['59', 'M9', 'ML', 'MINEURE', 'daily_drift_report vs build_xgb_training_set : 2 stratégies différentes'],
        ['60', 'M10', 'ML', 'MINEURE', '2 writers sur gold.model_drift_reports'],
        ['61', 'M11', 'ML', 'MINEURE', 'STGCNTrainer 6 horizons × VPS 12 Go RAM'],
        ['62', 'M12', 'ML', 'MINEURE', 'Pas de model_card pour Vélov et ST-GCN'],
        ['63', 'R3', 'Routing', 'MINEURE', 'Mauvais model_dir lookup (chemin mort)'],
        ['64', 'R4', 'Routing', 'MINEURE', '_compute_confidence = 0.85 hardcodé'],
        ['65', 'R5', 'Routing', 'MINEURE', 'channel_id retourné = f"{u}→{v}"'],
        ['66', 'R6', 'Routing', 'MINEURE', 'gold.tarifs_modes jamais lue par le code actif'],
        ['67', 'R7', 'Routing', 'MINEURE', 'TTL cache graph = 300 s hardcodé'],
        ['68', 'R8', 'Routing', 'MINEURE', 'Backend pathfinder rebuild du graphe incomplet'],
        ['69', 'R9', 'Routing', 'MINEURE', 'plan_transit_trip charge tous les lieux'],
        ['70', 'R10', 'Routing', 'MINEURE', '_get_current_day_type_and_bucket sans timezone'],
        ['71', 'G18', 'Gold', 'COSMETIQUE', '11 tables Gold legacy fantômes'],
        ['72', 'G19', 'Gold', 'COSMETIQUE', 'gold.infrastructure_bottlenecks.lat/lon synthétiques via HASHTEXT'],
    ],
    col_widths=[Cm(0.8), Cm(1.2), Cm(1.8), Cm(2.2), Cm(11)],
)

add_para(doc,
    "Total : 16 critiques, 19 majeures, 30 mineures, 2 cosmétiques (67 entrées — certaines couvrent plusieurs "
    "fichiers donc ~80 points d\'action uniques).", italic=True, color=C_MUTED)

page_break(doc)

# =========================================================
# 8. PLAN D'ACTION
# =========================================================

add_h1(doc, '8. Plan d\'action proposé — Sprint 16+')

add_para(doc, "Les actions sont triées par criticité et effort estimé. Aucune n\'est lancée sans validation explicite.")

add_h2(doc, '8.1 Sprint 16 (priorité 1) — Réconcilier schémas Silver/Gold')

add_para(doc, "Objectif : aligner le repo sur l\'état réel de la DB ou inversement.")
add_table(doc,
    headers=['#', 'Action', 'Fichiers', 'Effort'],
    rows=[
        ['1', 'Aligner DDL silver.trafic_boucles_clean : vitesse_kmh → déjà, ajouter measurement_time TIMESTAMPTZ, UNIQUE (channel_id, measurement_time). OU migrer le code INSERT vers speed_kmh.', 'init-db.sql:1279-1288 + bronze_to_silver.py:154-167', '1h'],
        ['2', 'Aligner DDL silver.meteo_hourly : ajouter temperature_c, rain_mm (ou migrer le code INSERT vers temperature_2m, precipitation).', 'init-db.sql:1258-1272 + bronze_to_silver.py:453', '2h'],
        ['3', 'Réconcilier silver.chantiers_actifs : choisir UN SEUL DDL entre trigger et GENERATED STORED, puis aligner le code INSERT.', 'init-db.sql:2425-2457 + bronze_to_silver.py:519-540', '3h'],
        ['4', 'Réconcilier gold.velov_features : aligner SQL et Python OU l\'inverse.', 'migrate_realign_v0.3.1.sql:125-146 + silver_to_gold.py:253-287 + xgboost_velov.py:25-37', '3h'],
    ],
    col_widths=[Cm(0.8), Cm(10), Cm(6.5), Cm(1.5)],
)

add_h2(doc, '8.2 Sprint 16 (priorité 2) — Réparer les INSERT cassés')

add_para(doc, "Objectif : éliminer les exceptions silencieuses.")
add_table(doc,
    headers=['#', 'Action', 'Fichiers', 'Effort'],
    rows=[
        ['5', 'Réparer _save_raw pour CalendrierScolaire + JoursFeries (colonnes fetched_at/raw_data absentes).', 'init-db.sql:108-115 + init-db.sql:382-385 + ingestion/calendrier_scolaire.py + ingestion/jours_feries.py', '2h'],
        ['6', 'Refacto base.py:145-155 pour re-raise les exceptions (fail loud).', 'src/ingestion/base.py', '1h'],
        ['7', 'Aligner migration_021_source_health.sql:130 : remplacer geom_wgs84 par geom.', 'scripts/sql/migration_021_source_health.sql', '15min'],
    ],
    col_widths=[Cm(0.8), Cm(10), Cm(6.5), Cm(1.5)],
)

add_h2(doc, '8.3 Sprint 16 (priorité 3) — Tables fantômes')

add_para(doc, "Objectif : faire le ménage.")
add_table(doc,
    headers=['#', 'Action', 'Fichiers', 'Effort'],
    rows=[
        ['8', 'Décider pour gold.velov_predictions, gold.amenagements_history, gold.fact_correlation_matrix, gold.mv_kpis_12_months : DROP les appels db_query ou créer les tables.', 'db_query.py:421, 833, 867, 888', '2h'],
        ['9', 'Peupler gold.channels_ref ou DROP la vue v_coherence_tomtom_vs_grandlyon.', 'init-db.sql:719-732 + migration_14_gold_coherence_tomtom_v2.sql', '2h'],
        ['10', 'DROP 11 tables Gold legacy fantômes (h3_trafic_live, h3_trafic_predictions, features_traffic, etc.).', 'init-db.sql (multiple lignes)', '1h'],
        ['11', 'DROP 6 tables Bronze orphelines (chantiers_historique, chantiers_voirie, comptages, parkings, prix_carburants, pvotrafic_snapshots).', 'init-db.sql (multiple lignes)', '1h'],
    ],
    col_widths=[Cm(0.8), Cm(10), Cm(6.5), Cm(1.5)],
)

add_h2(doc, '8.4 Sprint 17 (priorité 4) — Modèles ML et monitoring')

add_table(doc,
    headers=['#', 'Action', 'Fichiers', 'Effort'],
    rows=[
        ['12', 'Désactiver DAG retrain_xgboost_speed (DAG legacy hourly qui pollue avec 4 modèles .pkl obsolètes).', 'dags/ml/retrain_xgboost.py:101-115', '30min'],
        ['13', 'Propager MLFLOW_TRACKING_URI dans le code MLflowTracker + vérifier .env VPS.', 'mlflow_integration.py:98, 105 + .env', '1h'],
        ['14', 'Réparer MLflow register_model : artifact URI doit pointer vers un vrai artifact du run.', 'mlflow_integration.py:255-260', '1h'],
        ['15', 'Aligner model_version (1.0.0 vs 1.2.0) entre train et predict.', 'xgboost_speed.py:208 vs 334', '30min'],
        ['16', 'Ajouter UNIQUE INDEX (axis_key, calculated_at) sur gold.mv_xgb_vs_tomtom pour permettre REFRESH CONCURRENTLY.', 'migration_020_xgb_vs_tomtom.sql:99-104', '30min'],
        ['17', 'Décider du sort du ST-GCN : l\'activer, le supprimer, ou le mettre en sandbox.', 'retrain_gnn.py:310 + src/ml/model_registry.py:200-244', '2h discussion'],
    ],
    col_widths=[Cm(0.8), Cm(10), Cm(6.5), Cm(1.5)],
)

add_h2(doc, '8.5 Sprint 17+ — Dettes de long terme')

add_table(doc,
    headers=['#', 'Action', 'Fichiers', 'Effort'],
    rows=[
        ['18', 'Faire respecter horizon_minutes dans le pathfinder (routage H+1h réel).', 'src/routing/graph.py:290-298', '1 sem'],
        ['19', 'Refacto pour utiliser gold.tarifs_modes au lieu des constantes hardcodées.', 'eco_calculator.py:50-66 + migration_016', '3h'],
        ['20', 'Implémenter recommend_mode avec scoring composite 50/30/20 comme spec l\'annonce.', 'eco_calculator.py:225-281 + mode_comparison.py', '4h'],
        ['21', 'Ajouter tests unitaires sur bronze_to_silver.py (547 lignes sans test).', 'tests/unit/test_bronze_to_silver.py (à créer)', '3h'],
        ['22', 'Brancher les compteurs Prometheus sur les collecteurs Bronze (n_requests, n_failures, last_success_at).', 'src/ingestion/base.py + src/api/metrics.py', '2h'],
    ],
    col_widths=[Cm(0.8), Cm(10), Cm(6.5), Cm(1.5)],
)

page_break(doc)

# =========================================================
# 9. ANNEXES
# =========================================================

add_h1(doc, '9. Annexes')

add_h2(doc, '9.1 Liste des fichiers audités')

add_para(doc, "Collecte / Bronze :")
add_mono(doc, """
  dags/bronze/collect_bronze.py
  dags/bronze/collect_calendriers_monthly.py
  dags/bronze/collect_tomtom_traffic.py
  src/ingestion/__init__.py
  src/ingestion/base.py
  src/ingestion/trafic_grandlyon.py
  src/ingestion/tcl_siri_lite.py
  src/ingestion/velov.py
  src/ingestion/meteo.py
  src/ingestion/air_quality.py
  src/ingestion/chantiers.py
  src/ingestion/calendrier_scolaire.py
  src/ingestion/jours_feries.py
  src/ingestion/tomtom_traffic.py
""")

add_para(doc, "Transformations :")
add_mono(doc, """
  dags/transforms/transform_bronze_to_silver.py
  dags/transforms/transform_silver_to_gold.py
  dags/transforms/build_spatial_mapping.py
  src/transformation/bronze_to_silver.py
  src/transformation/silver_to_gold.py
""")

add_para(doc, "Couche data (consommateurs) :")
add_mono(doc, """
  src/data/db_query.py
  src/data/data_loader.py
  src/data/airflow_client.py
  src/data/labels.py
  src/data/tcl_lines.py
""")

add_para(doc, "Modèles ML :")
add_mono(doc, """
  src/models/xgboost_speed.py
  src/models/xgboost_velov.py
  src/models/stgcn_wrapper.py
  training/stgcn/dataset.py
  training/stgcn/model.py
  training/stgcn/train.py
  dags/ml/dag_daily_speed_train.py
  dags/ml/dag_inference_xgboost.py
  dags/ml/build_xgb_training_set.py
  dags/ml/refresh_xgb_vs_tomtom.py
  dags/ml/retrain_gnn.py
  dags/ml/retrain_xgboost.py
  dags/ml/daily_drift_report.py
""")

add_para(doc, "Routing :")
add_mono(doc, """
  src/routing/pathfinder_multimodal.py
  src/routing/pathfinder.py
  src/routing/graph.py
  src/routing/eco_calculator.py
  src/routing/snap_to_roads.py
""")

add_para(doc, "Migrations SQL :")
add_mono(doc, """
  deploy/init-db.sql
  scripts/migrate_realign_v0.3.1.sql
  scripts/sql/audit_dim_spatial_writers.sql
  scripts/sql/backfill_dim_spatial_lat_lon.sql
  scripts/sql/create_lieux_calendrier.sql
  scripts/sql/create_lieux_velov_proches.sql
  scripts/sql/create_mv_line_kpis_otp.sql
  scripts/sql/create_mv_twgid_to_lyo.sql
  scripts/sql/create_pathfinder_helpers.sql
  scripts/sql/create_referentiel_lieux.sql
  scripts/sql/create_referentiel_transports.sql
  scripts/sql/create_tomtom_traffic.sql
  scripts/sql/create_velov_maillage.sql
  scripts/sql/create_xgb_training_set.sql
  scripts/sql/migration_14_gold_coherence_tomtom_v2.sql
  scripts/sql/migration_15_aggregate_line_ref.sql
  scripts/sql/migration_016_tarifs_modes.sql
  scripts/sql/migration_017_multimodal_grid.sql
  scripts/sql/migration_018_bus_traffic_spatial.sql
  scripts/sql/migration_019_network_health.sql
  scripts/sql/migration_020_xgb_vs_tomtom.sql
  scripts/sql/migration_021_source_health.sql
""")

add_h2(doc, '9.2 Variables d\'environnement critiques')

add_table(doc,
    headers=['Variable', 'Consommateurs', 'Valeur par défaut'],
    rows=[
        ['POSTGRES_PASSWORD', 'Toute la couche data (db_query, data_loader)', 'aucune (requis)'],
        ['MLFLOW_TRACKING_URI', 'MLflowTracker', 'http://localhost:5000 ⚠'],
        ['TOMTOM_API_KEY', 'TomTomTrafficFlow', 'aucune (DAG no-op si absent)'],
        ['GRANDLYON_USERNAME / PASSWORD', '4 collecteurs Grand Lyon', 'aucune (auth optionnelle WFS)'],
        ['LYON_LATITUDE / LYON_LONGITUDE', 'meteo, air_quality, chantiers', '45.7640 / 4.8357'],
        ['LYONFLOW_DEMO_MODE', 'helpers mock', '0 (politique zéro mock)'],
        ['AIRFLOW_FERNET_KEY', 'Airflow', 'aucune (requis)'],
        ['LYONFLOW_API_KEY', 'FastAPI', 'aucune (requis)'],
    ],
    col_widths=[Cm(5), Cm(8), Cm(4.5)],
)

add_h2(doc, '9.3 Notes méthodologiques')

add_para(doc,
    "Cet audit s\'appuie exclusivement sur la lecture du code source à la date du 2026-06-20. "
    "Aucun mock, aucune supposition, aucune interpolation. Chaque doute est sourcé par un fichier:ligne."
)
add_para(doc,
    "Les \"hypothèses hautes\" mentionnées dans le rapport (notamment S2, S3, S4, G1) reposent sur l\'hypothèse "
    "qu\'un ALTER TABLE non commité a été appliqué sur le VPS pour aligner les schémas effective ↔ code. "
    "Cette pratique est documentée dans AGENTS.md (« patch in-place non commité ») et explique pourquoi la "
    "production fonctionne malgré les incohérences du repo."
)
add_para(doc,
    "Recommandation forte : à l\'avenir, toute modification de schéma doit être committée dans scripts/sql/ "
    "AVANT d\'être appliquée sur la DB, pour que le repo reste la source de vérité unique."
)

# Final
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Fin du document —')
r.font.size = Pt(10)
r.italic = True
r.font.color.rgb = C_MUTED


# Save
out = '/Users/patriceduclos/Documents/Lyonfull/docs/AUDIT_PIPELINE_DATA_2026-06-20.docx'
doc.save(out)
print(f"OK: {out}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
